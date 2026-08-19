"""把網站內的護腎知識輸出成 Word 檔（審閱用）。"""
import json, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

DATA = json.load(open("knowledge_export.json", encoding="utf-8"))
CATS = ["血壓管理", "血糖管理", "血脂代謝", "用藥安全", "飲食護腎", "生活習慣", "檢查數值", "警訊與迷思"]
ACCENT = RGBColor(0x7C, 0x5C, 0xFF)
INK    = RGBColor(0x3D, 0x2C, 0x29)
DIM    = RGBColor(0x8A, 0x72, 0x68)
FONT   = "Microsoft JhengHei"

doc = Document()

# ── 版面與預設字型 ──
sec = doc.sections[0]
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, attr, Cm(2.2))

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def run(p, text, size=11, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color is not None:
        r.font.color.rgb = color
    return r


def bottom_border(p, color="7C5CFF", size=12):
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def break_before(p):
    """在段落前分頁 — 比插入分頁符段落乾淨，不會留下空段落造成空白頁。"""
    p.paragraph_format.page_break_before = True
    return p


# ── 封面 ──
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "護腎教室", size=36, bold=True, color=ACCENT)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "KidneyGod.Studio", size=13, color=DIM)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "護腎知識全集", size=22, bold=True, color=INK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, f"共 {len(DATA)} 篇．八大分類", size=12, color=DIM)
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "審閱用文稿　2026 年 8 月 19 日", size=11, color=DIM)

# ── 目錄 ──
p = break_before(doc.add_paragraph()); bottom_border(p)
run(p, "目錄", size=16, bold=True, color=ACCENT)
doc.add_paragraph()
idx = 0
for cat in CATS:
    lst = [d for d in DATA if d["cat"] == cat]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, f"{cat}（{len(lst)} 篇）", size=12, bold=True, color=INK)
    for d in lst:
        idx += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(0)
        run(p, f"{idx}. {d['title']}", size=10.5, color=DIM)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 內容 ──
n = 0
for ci, cat in enumerate(CATS):
    lst = [d for d in DATA if d["cat"] == cat]
    p = break_before(doc.add_paragraph()); bottom_border(p)
    run(p, cat, size=16, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    run(p, f"本類共 {len(lst)} 篇", size=9.5, color=DIM)

    for d in lst:
        n += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        run(p, f"{n}. {d['title']}", size=13, bold=True, color=INK)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run(p, f"商品：【{d['brand']}】{d['emoji']}　｜　售價 {d['price']} 腎元　｜　代碼 {d['id']}",
            size=9, color=DIM, italic=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        run(p, d["body"], size=11)

# ── 免責聲明 ──
p = break_before(doc.add_paragraph()); bottom_border(p)
run(p, "醫療免責聲明", size=16, bold=True, color=ACCENT)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(8)
run(p, "本文件內容為一般性健康衛教知識，不能取代醫師診斷與治療建議。"
       "腎功能異常、用藥問題請諮詢腎臟科醫師或藥師；有症狀請就醫。", size=11)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run(p, "內容出處：護腎教室 KidneyGod.Studio"
       "（https://kidneygod-studio.github.io/kidneygod/）之遊戲內知識商品，全數為原創撰寫。",
    size=9.5, color=DIM)

doc.core_properties.title = "護腎知識全集"
doc.core_properties.author = "護腎教室 KidneyGod.Studio"

doc.save("護腎知識全集.docx")
print(f"OK - {n} articles written")

# ── 同時輸出純文字版 ──
lines = ["護腎教室 KidneyGod.Studio — 護腎知識全集",
         f"共 {len(DATA)} 篇．八大分類",
         "=" * 60, ""]
n = 0
for cat in CATS:
    lst = [d for d in DATA if d["cat"] == cat]
    lines += [f"■ {cat}（{len(lst)} 篇）", "-" * 60, ""]
    for d in lst:
        n += 1
        lines += [f"{n}. {d['title']}",
                  f"   商品：【{d['brand']}】{d['emoji']} | 售價 {d['price']} 腎元 | 代碼 {d['id']}",
                  "",
                  f"   {d['body']}",
                  ""]
    lines.append("")
lines += ["=" * 60,
          "【醫療免責聲明】",
          "本文件內容為一般性健康衛教知識，不能取代醫師診斷與治療建議。",
          "腎功能異常、用藥問題請諮詢腎臟科醫師或藥師；有症狀請就醫。",
          "",
          "內容出處：護腎教室 KidneyGod.Studio",
          "https://kidneygod-studio.github.io/kidneygod/"]
open("護腎知識全集.txt", "w", encoding="utf-8").write("\n".join(lines))
print("TXT written")
